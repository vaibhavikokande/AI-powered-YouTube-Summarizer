import io
from dataclasses import dataclass
from xml.sax.saxutils import escape

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.core.exceptions import ValidationAppError
from app.models.summary import Summary

_SUPPORTED_FORMATS = {"pdf", "docx", "markdown", "txt"}


@dataclass
class ExportedFile:
    content: bytes
    media_type: str
    filename: str


class ExportService:
    """Renders a Summary into a downloadable file: PDF, DOCX, Markdown, or plain text.

    Each format is built directly from the Summary's fields rather than by
    converting one format's output into another (e.g. stripping markdown
    syntax out of the Markdown export to get a text export) — LLM-generated
    prose can legitimately contain characters like "- " or "> " mid-sentence,
    so a naive string-replace pass would corrupt content, not just formatting.
    """

    def export(self, summary: Summary, video_title: str | None, export_format: str) -> ExportedFile:
        if export_format not in _SUPPORTED_FORMATS:
            raise ValidationAppError(
                f"Unsupported export format {export_format!r}; "
                f"choose one of {sorted(_SUPPORTED_FORMATS)}."
            )

        base_name = self._safe_filename(video_title or "summary")

        if export_format == "txt":
            return ExportedFile(self._to_txt(summary, video_title), "text/plain", f"{base_name}.txt")
        if export_format == "markdown":
            return ExportedFile(
                self._to_markdown(summary, video_title), "text/markdown", f"{base_name}.md"
            )
        if export_format == "docx":
            return ExportedFile(
                self._to_docx(summary, video_title),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                f"{base_name}.docx",
            )
        return ExportedFile(self._to_pdf(summary, video_title), "application/pdf", f"{base_name}.pdf")

    @staticmethod
    def _safe_filename(title: str) -> str:
        cleaned = "".join(c if c.isalnum() or c in " -_" else "" for c in title).strip()
        cleaned = cleaned.replace(" ", "_")[:80]
        return cleaned or "summary"

    @staticmethod
    def _format_timestamp(seconds: int) -> str:
        minutes, secs = divmod(int(seconds), 60)
        hours, minutes = divmod(minutes, 60)
        if hours:
            return f"{hours}:{minutes:02d}:{secs:02d}"
        return f"{minutes}:{secs:02d}"

    def _to_markdown(self, summary: Summary, video_title: str | None) -> bytes:
        title = video_title or "Video Summary"
        lines = [f"# {title}", "", f"**Summary type:** {summary.summary_type.value}", "", summary.content, ""]

        if summary.topics.get("main_topics"):
            lines.append("## Topics")
            lines.extend(f"- {t}" for t in summary.topics["main_topics"])
            lines.append("")

        takeaways = summary.key_takeaways
        if takeaways.get("important_concepts"):
            lines.append("## Key Concepts")
            lines.extend(f"- {c}" for c in takeaways["important_concepts"])
            lines.append("")
        if takeaways.get("action_items"):
            lines.append("## Action Items")
            lines.extend(f"- {a}" for a in takeaways["action_items"])
            lines.append("")
        if takeaways.get("important_quotes"):
            lines.append("## Quotes")
            lines.extend(f"> {q}" for q in takeaways["important_quotes"])
            lines.append("")
        if takeaways.get("statistics"):
            lines.append("## Statistics")
            lines.extend(f"- {s}" for s in takeaways["statistics"])
            lines.append("")

        if summary.timestamped_sections:
            lines.append("## Timestamped Sections")
            for section in summary.timestamped_sections:
                ts = self._format_timestamp(section["timestamp_seconds"])
                lines.append(f"- **[{ts}] {section['title']}** — {section['summary']}")
            lines.append("")

        if summary.mindmap_markdown:
            lines.append("## Mind Map")
            lines.append(summary.mindmap_markdown)
            lines.append("")

        return "\n".join(lines).encode("utf-8")

    def _to_txt(self, summary: Summary, video_title: str | None) -> bytes:
        title = video_title or "Video Summary"
        lines = [title, "=" * len(title), "", f"Summary type: {summary.summary_type.value}", "", summary.content, ""]

        if summary.topics.get("main_topics"):
            lines.append("TOPICS")
            lines.extend(f"- {t}" for t in summary.topics["main_topics"])
            lines.append("")

        takeaways = summary.key_takeaways
        if takeaways.get("important_concepts"):
            lines.append("KEY CONCEPTS")
            lines.extend(f"- {c}" for c in takeaways["important_concepts"])
            lines.append("")
        if takeaways.get("action_items"):
            lines.append("ACTION ITEMS")
            lines.extend(f"- {a}" for a in takeaways["action_items"])
            lines.append("")
        if takeaways.get("important_quotes"):
            lines.append("QUOTES")
            lines.extend(f'"{q}"' for q in takeaways["important_quotes"])
            lines.append("")
        if takeaways.get("statistics"):
            lines.append("STATISTICS")
            lines.extend(f"- {s}" for s in takeaways["statistics"])
            lines.append("")

        if summary.timestamped_sections:
            lines.append("TIMESTAMPED SECTIONS")
            for section in summary.timestamped_sections:
                ts = self._format_timestamp(section["timestamp_seconds"])
                lines.append(f"[{ts}] {section['title']} - {section['summary']}")
            lines.append("")

        return "\n".join(lines).encode("utf-8")

    def _to_docx(self, summary: Summary, video_title: str | None) -> bytes:
        document = Document()
        document.add_heading(video_title or "Video Summary", level=1)
        document.add_paragraph(f"Summary type: {summary.summary_type.value}")
        document.add_paragraph(summary.content)

        if summary.topics.get("main_topics"):
            document.add_heading("Topics", level=2)
            for t in summary.topics["main_topics"]:
                document.add_paragraph(t, style="List Bullet")

        takeaways = summary.key_takeaways
        if takeaways.get("important_concepts"):
            document.add_heading("Key Concepts", level=2)
            for c in takeaways["important_concepts"]:
                document.add_paragraph(c, style="List Bullet")
        if takeaways.get("action_items"):
            document.add_heading("Action Items", level=2)
            for a in takeaways["action_items"]:
                document.add_paragraph(a, style="List Bullet")
        if takeaways.get("important_quotes"):
            document.add_heading("Quotes", level=2)
            for q in takeaways["important_quotes"]:
                document.add_paragraph(q, style="Intense Quote")
        if takeaways.get("statistics"):
            document.add_heading("Statistics", level=2)
            for s in takeaways["statistics"]:
                document.add_paragraph(s, style="List Bullet")

        if summary.timestamped_sections:
            document.add_heading("Timestamped Sections", level=2)
            for section in summary.timestamped_sections:
                ts = self._format_timestamp(section["timestamp_seconds"])
                document.add_paragraph(f"[{ts}] {section['title']}: {section['summary']}")

        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _to_pdf(self, summary: Summary, video_title: str | None) -> bytes:
        # reportlab's Paragraph treats its input as markup, so anything from
        # an LLM (which can legitimately contain "&", "<", ">") must be
        # XML-escaped first or it can break the parser or render incorrectly.
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=LETTER)
        styles = getSampleStyleSheet()

        story = [Paragraph(escape(video_title or "Video Summary"), styles["Title"]), Spacer(1, 12)]
        story.append(Paragraph(f"Summary type: {summary.summary_type.value}", styles["Normal"]))
        story.append(Spacer(1, 12))
        story.append(Paragraph(escape(summary.content).replace("\n", "<br/>"), styles["BodyText"]))
        story.append(Spacer(1, 12))

        if summary.topics.get("main_topics"):
            story.append(Paragraph("Topics", styles["Heading2"]))
            for t in summary.topics["main_topics"]:
                story.append(Paragraph(f"• {escape(t)}", styles["Normal"]))
            story.append(Spacer(1, 12))

        takeaways = summary.key_takeaways
        if takeaways.get("important_concepts"):
            story.append(Paragraph("Key Concepts", styles["Heading2"]))
            for c in takeaways["important_concepts"]:
                story.append(Paragraph(f"• {escape(c)}", styles["Normal"]))
            story.append(Spacer(1, 12))
        if takeaways.get("action_items"):
            story.append(Paragraph("Action Items", styles["Heading2"]))
            for a in takeaways["action_items"]:
                story.append(Paragraph(f"• {escape(a)}", styles["Normal"]))
            story.append(Spacer(1, 12))

        if summary.timestamped_sections:
            story.append(Paragraph("Timestamped Sections", styles["Heading2"]))
            for section in summary.timestamped_sections:
                ts = self._format_timestamp(section["timestamp_seconds"])
                title = escape(section["title"])
                section_summary = escape(section["summary"])
                story.append(
                    Paragraph(f"[{ts}] <b>{title}</b> — {section_summary}", styles["Normal"])
                )

        doc.build(story)
        return buffer.getvalue()
